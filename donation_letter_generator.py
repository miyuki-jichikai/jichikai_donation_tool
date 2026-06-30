# ============================================================
# 寄付依頼書 自動生成ツール
# donation_letter_generator.py  v2.7
# ============================================================
# 変更点 (v2.6 → v2.7):
#   - EXE化(--onefile)時にprogress.jsonの保存先が毎回変わってしまうバグを修正
#     （__file__がPyInstallerの一時展開フォルダを指していたため。
#      sys.executableを基準にすることでEXE自体のフォルダを正しく取得）
#   - 広告選択画面にソフトウェアバージョンを表示（タイトルバー＋画面右上）
# 変更点 (v2.5 → v2.6):
#   - f文字列破損バグ修正（len(completed, parent=_tk_root)のTypeError解消）
# 変更点 (v2.4 → v2.5):
#   - EXE化後に再開時の広告選択GUIが反応しなくなるバグを修正
#   - tk.Tk()は1回だけ作成し、各ダイアログ/GUIはtk.Toplevel()を使用するよう変更
# 変更点 (v2.3 → v2.4):
#   - Excelファイル名を260831_プログラム広告掲載リストTEST.xlsxに変更
#   - Excel列構成の変更に対応（A列IDで種別判定・G列寄付金額・X列ページ番号）
#   - 抽出条件変更：IDに「内」または「外」を含む行を処理対象とする
#   - 起動時にExcel・PDF・Wordテンプレートのファイル選択ダイアログを追加
#     （Excel:.xlsxのみ / PDF:.pdfのみ / テンプレート:.dotxのみ）
#   - 広告選択GUIの縦サイズを縮小（ノートPC対応）
# 変更点 (v2.2 → v2.3):
#   - F列のみによるシンプルなページ決定に戻す
#   - ページ番号直接入力ジャンプ追加
#   - 自動改ページ（下端＆右端30mm判定）
#   - 再開時ページ復元
#   - bm_numberに会社名＋様・金額追加
# 変更点 (v2.1 → v2.2):
#   - ページ検索を「会社名 OR 住所（G列）」の両方で実施
#   - 会社名がロゴ画像の場合でも住所テキストでページを特定
#   - スコアリング導入：住所一致(+3) > 会社名一致(+2) で優先
#   - 住所から町名を自動抽出して検索精度を向上
#   - コンソールに検索ヒット理由を表示
# 変更点 (v2.0 → v2.1):
#   - F列（プログラムページ）の手動入力が不要になった
#   - PDFのテキストから会社名を自動検索して該当ページへジャンプ
#   - GUIにページ送りボタン（◀前 / 次▶）を追加
#   - 自動検索が外れた場合もページを手動で移動可能
#
# 必要ライブラリ:
#   pip install pillow python-docx openpyxl pdf2image pdfplumber
#   poppler (PATH設定済みであること)
#
# 使い方:
#   python donation_letter_generator.py
# ============================================================
 
import os
import sys
import json
import unicodedata
import shutil
import tempfile
import zipfile
import tkinter as tk
from tkinter import messagebox, filedialog
from PIL import Image, ImageTk
import pandas as pd
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
 
try:
    import pdfplumber
except ImportError:
    print("❌ pdfplumber がインストールされていません。")
    print("   pip install pdfplumber を実行してください。")
    sys.exit(1)
 
# ============================================================
# ★★★ 設定 ★★★
# ============================================================
def get_base_dir():
    """
    実行ファイルの実体があるフォルダを返す。
    - 通常の .py 実行時 : スクリプトファイルのあるフォルダ
    - PyInstaller --onefile でEXE化した場合:
        __file__ は一時解凍フォルダ(sys._MEIPASS, 毎回変わる)を指してしまうため、
        sys.executable（EXE自体のパス）を基準にする。
    """
    if getattr(sys, 'frozen', False):
        # PyInstallerでビルドされたEXEとして実行されている
        return os.path.dirname(os.path.abspath(sys.executable))
    else:
        # 通常の python スクリプトとして実行されている
        return os.path.dirname(os.path.abspath(__file__))


APP_VERSION   = "v2.7"

BASE_DIR      = get_base_dir()
DATA_DIR      = os.path.join(BASE_DIR, "data")
OUTPUT_DIR    = os.path.join(BASE_DIR, "output")
TMP_DIR       = os.path.join(BASE_DIR, "tmp_ad_images")
 
EXCEL_FILE    = os.path.join(DATA_DIR, "260831_プログラム広告掲載リストTEST.xlsx")
PDF_FILE      = os.path.join(DATA_DIR, "ad_list.pdf")
TEMPLATE_FILE = os.path.join(DATA_DIR, "企業寄付申込書.dotx")
PROGRESS_FILE = os.path.join(BASE_DIR, "progress.json")
 
# popplerのパス
POPPLER_PATH  = r"C:\poppler\Library\bin"
 
# PDFの広告ページが何ページ目から始まるか（表紙・目次などを除いた開始ページ）
PDF_AD_START_PAGE = 6  # ← このPDFは6ページ目から広告が始まる
 
# 寄付金額ごとの広告サイズ (幅mm, 高さmm)
AD_SIZES = {
    100000: (160, 240),
     50000: (160, 240),
     30000: (150, 115),
     20000: (150,  75),
     10000: (150,  45),
      5000: ( 75,  45),
      3000: ( 75,  35),
}
 
# ============================================================
# ユーティリティ
# ============================================================
 
def normalize_str(s):
    """全角→半角・スペース除去・小文字化して比較しやすくする"""
    s = unicodedata.normalize('NFKC', str(s))
    return s.replace(' ', '').replace('　', '').lower().strip()
 
 
def get_ad_size(amount):
    """寄付金額から広告サイズ(幅mm, 高さmm)を返す"""
    for threshold in sorted(AD_SIZES.keys(), reverse=True):
        if amount >= threshold:
            return AD_SIZES[threshold]
    return AD_SIZES[3000]
 
 
def load_progress():
    """進捗ファイルを読み込む"""
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"completed_ids": [], "last_index": 0}
 
 
def save_progress(progress):
    """進捗ファイルに保存する"""
    with open(PROGRESS_FILE, 'w', encoding='utf-8') as f:
        json.dump(progress, f, ensure_ascii=False, indent=2)
    print(f"  → 進捗を保存しました")
 
 
def load_excel(excel_path=None):
    """Excelから処理対象行を読み込む（新列構成版）
 
    新しい列構成:
      A列(0) : ID         ← 「内xx」=町内寄付 / 「外xx」=町外寄付 で種別判定
      B列(1) : 区部
      C列(2) : 会社・店名・氏名
      D列(3) : 住所
      E列(4) : 電話番号
      F列(5) : 担当者
      G列(6) : 25年寄付金  ← 寄付金額（列位置固定）
      I列(8) : ふりがな
      J列(9) : 昨年と同じか
      K列(10): 記載事項
      X列(23): プログラムページ
    """
    df = pd.read_excel(excel_path or EXCEL_FILE, header=0, dtype=str)
 
    # ── G列（7列目・0始まりで6番目）の寄付金額を列位置で取得 ──
    g_col_name = df.columns[6]   # G列
    print(f"  → 寄付金額(G列): 「{g_col_name}」を使用")
    df['_金額'] = pd.to_numeric(
        df[g_col_name], errors='coerce').fillna(0).astype(int)
    df['_25寄付金'] = df['_金額'].copy()
 
    # ── X列（24列目・0始まりで23番目）のプログラムページを取得 ──
    if len(df.columns) > 23:
        x_col_name = df.columns[23]   # X列
        print(f"  → プログラムページ(X列): 「{x_col_name}」を使用")
        df['_page'] = pd.to_numeric(
            df[x_col_name], errors='coerce').fillna(0).astype(int)
    else:
        df['_page'] = 0
        print("  ⚠ X列「プログラムページ」が見つかりません。全行0として扱います。")
 
    # ── A列のIDに「内」または「外」が含まれる行を抽出 ──
    # 「内」→ 町内寄付 / 「外」→ 町外寄付 に相当
    id_col = df.columns[0]   # A列
    mask = df[id_col].str.contains('内|外', na=False)
    target = df[mask].reset_index(drop=True)
    print(f"  → 処理対象件数: {len(target)} 件"
          f"（IDに「内」または「外」を含む行）")
    return target
 
 
# ============================================================
# PDF処理：画像変換 ＋ テキスト抽出
# ============================================================
 
def pdf_to_images(pdf_path):
    """PDFを全ページ画像に変換して {ページ番号: PIL.Image} を返す"""
    print("PDFを画像に変換中（しばらくお待ちください）...")
    images = convert_from_path(
        pdf_path, dpi=150, poppler_path=POPPLER_PATH)
    page_map = {}
    for i, img in enumerate(images):
        page_num = PDF_AD_START_PAGE + i
        page_map[page_num] = img
    print(f"  → {len(page_map)} ページ変換完了"
          f" (ページ {min(page_map)}〜{max(page_map)})")
    return page_map
 
 
def extract_pdf_texts(pdf_path):
    """
    pdfplumber でPDF全ページのテキストを抽出する。
    戻り値: {ページ番号: テキスト文字列}
    ページ番号は pdf_to_images と同じ基準（PDF_AD_START_PAGE始まり）。
    """
    print("PDFのテキストを抽出中...")
    page_texts = {}
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_num = PDF_AD_START_PAGE + i
            text = page.extract_text() or ""
            page_texts[page_num] = normalize_str(text)
    print(f"  → {len(page_texts)} ページのテキスト抽出完了")
    return page_texts
 
 
def parse_address(address):
    """
    住所文字列から「町名」と「番地」を分けて返す。
 
    例：「中央区三幸町476-8」  → town="三幸町"   banchi="476-8"
    例：「浜名区都田町7555-47」 → town="都田町"   banchi="7555-47"
    例：「中央区豊岡429-9」    → town="豊岡"     banchi="429-9"
      ※「豊岡」は[町丁区村字目]で終わらないが、番地前の最後の漢字列を使う
 
    戻り値: (town: str, banchi: str)
      どちらも正規化済み。取得できない場合は空文字。
    """
    import re
    s = unicodedata.normalize('NFKC', str(address)).strip()
 
    # ── 番地を抽出
    #    「数字 (ハイフン|数字)+ 」の形を番地とみなす
    #    例: 476-8 / 7555-47 / 429-9 / 1823-2 / 255 など
    banchi = ''
    banchi_match = re.search(r'(\d[\d\-－]+)', s)
    if banchi_match:
        raw = banchi_match.group(1).rstrip('-－')
        if len(raw) >= 2:
            banchi = normalize_str(raw)
 
    # ── 町名を抽出（番地より前の部分から）
    before_banchi = s[:banchi_match.start()] if banchi_match else s
 
    town = ''
    # ① 「町・丁・区・村・字・目」で終わる漢字列（最長）
    town_matches = re.findall(
        r'[\u4e00-\u9fff\u30a0-\u30ff\u3040-\u309f\d]+[町丁区村字目]',
        before_banchi)
    if town_matches:
        town = normalize_str(max(town_matches, key=len))
    else:
        # ② 上記で取れない場合（「豊岡」など）→ 番地直前の漢字列を使う
        kanji_matches = re.findall(
            r'[\u4e00-\u9fff]+', before_banchi)
        if kanji_matches:
            town = normalize_str(kanji_matches[-1])  # 最後の漢字列
 
    return town, banchi
 
 
def _score_name(query, page_text, min_len=3, weight=2):
    """
    会社名スコア：部分一致（最長一致）× weight。
    """
    if not query or not page_text:
        return 0
    if query in page_text:
        return len(query) * weight
    best = 0
    for start in range(len(query)):
        for end in range(start + min_len, len(query) + 1):
            part = query[start:end]
            if part in page_text and len(part) > best:
                best = len(part)
    return best * weight
 
 
def _score_banchi(banchi, page_text, weight=20):
    """
    番地スコア：完全一致のみ・高得点。
    部分一致は使わない（隣のページの番地との誤マッチ防止）。
 
    「完全一致」の定義：
      番地の前後が数字・ハイフン以外（スペース・改行・漢字など）であること。
      例: "476-8" が "1476-8" の一部にマッチしないよう境界チェックする。
    """
    import re
    if not banchi or not page_text:
        return 0
    # 番地の前後に数字・ハイフンが来ない位置でのみマッチ
    pattern = r'(?<![0-9\-])' + re.escape(banchi) + r'(?![0-9\-])'
    if re.search(pattern, page_text):
        return len(banchi) * weight
    return 0
 
 
def _score_town(town, page_text, weight=3):
    """
    町名スコア：部分一致（最長一致）× weight。
    補助的な役割なので重みは低め。
    """
    if not town or not page_text:
        return 0
    if town in page_text:
        return len(town) * weight
    best = 0
    for start in range(len(town)):
        for end in range(start + 2, len(town) + 1):
            part = town[start:end]
            if part in page_text and len(part) > best:
                best = len(part)
    return best * weight
 
 
def normalize_tel(tel):
    """
    電話番号を正規化して検索しやすい形に変換する。
    ・全角→半角
    ・ハイフン・スペース・括弧などを除去
    ・先頭の「0」は残す（市外局番）
    例: 「053-421-1234」→「0534211234」
        「(053)421-1234」→「0534211234」
    """
    s = unicodedata.normalize('NFKC', str(tel)).strip()
    import re
    s = re.sub(r'[^\d]', '', s)   # 数字以外を除去
    return s
 
 
def _score_tel(tel_normalized, page_text, weight=20):
    """
    電話番号スコア：正規化した番号が正規化したページテキストに
    完全一致（境界チェック付き）した場合のみ高得点。
    番地と同等の重み。
    """
    import re
    if not tel_normalized or len(tel_normalized) < 6:
        return 0
    # ページテキストも数字のみに正規化して比較
    page_nums = re.sub(r'[^\d]', '', page_text)
    if tel_normalized in page_nums:
        return len(tel_normalized) * weight
    return 0
 
 
def find_page_by_company_and_address(
        company_name, address, tel, page_texts, prev_page=None):
    """
    会社名・番地・電話番号（完全一致）・町名の4軸でスコアリングし、
    最も一致するページ番号を返す。見つからない場合は None。
 
    重み付け:
      電話番号 × 20  ← 完全一致のみ・最もユニーク
      番地     × 20  ← 完全一致のみ・境界チェック付き
      会社名   ×  2  ← 部分一致（ロゴ画像の場合はヒットしないことがある）
      町名     ×  3  ← 部分一致・補助的
 
    順番ヒント (prev_page):
      広告はExcelの会社順に並んでいるため、
      電話番号・番地・会社名がどちらもヒットせず町名スコアのみの場合は
      prev_page 以降のページに絞って再検索する（誤マッチ防止）。
    """
    name_query         = normalize_str(company_name)
    town_query, banchi = parse_address(address)
    tel_normalized     = normalize_tel(tel)
 
    page_scores = {}  # {page_num: (total, name_s, banchi_s, tel_s, town_s)}
    for page_num, text in page_texts.items():
        name_s   = _score_name(name_query,    text)
        banchi_s = _score_banchi(banchi,      text)
        tel_s    = _score_tel(tel_normalized, text)
        town_s   = _score_town(town_query,    text)
        total    = name_s + banchi_s + tel_s + town_s
        if total > 0:
            page_scores[page_num] = (total, name_s, banchi_s, tel_s, town_s)
 
    if not page_scores:
        return None, "ヒットなし"
 
    # ── 電話番号・番地・会社名のいずれかがヒット → 「強ヒット」
    strong = {p: v for p, v in page_scores.items()
              if v[1] > 0 or v[2] > 0 or v[3] > 0}  # name/banchi/tel
 
    if strong:
        if prev_page is not None:
            strong_after = {p: v for p, v in strong.items()
                            if p >= prev_page}
            candidates = strong_after if strong_after else strong
        else:
            candidates = strong
        best_page = max(candidates, key=lambda p: candidates[p][0])
        reason_suffix = ""
    else:
        # 町名スコアのみ → prev_page 以降に絞って誤マッチを減らす
        if prev_page is not None:
            after = {p: v for p, v in page_scores.items()
                     if p >= prev_page}
            candidates = after if after else page_scores
            reason_suffix = "（順番ヒント使用）"
        else:
            candidates = page_scores
            reason_suffix = ""
        best_page = max(candidates, key=lambda p: candidates[p][0])
 
    total, name_s, banchi_s, tel_s, town_s = page_scores[best_page]
 
    reasons = []
    if name_s   > 0: reasons.append(f"会社名:{name_s}")
    if banchi_s > 0: reasons.append(f"番地「{banchi}」:{banchi_s}")
    if tel_s    > 0: reasons.append(f"電話「{tel_normalized}」:{tel_s}")
    if town_s   > 0: reasons.append(f"町名「{town_query}」:{town_s}")
    reason = " / ".join(reasons) + reason_suffix
 
    return best_page, reason
 
 
# ============================================================
# Word出力
# ============================================================
 
def open_template(template_path=None):
    """.dotxをdocxとして開く（ZIPのcontent_typeを書き換え）"""
    tmp_tpl = tempfile.mktemp(suffix='.docx')
    shutil.copy(template_path or TEMPLATE_FILE, tmp_tpl)
    with zipfile.ZipFile(tmp_tpl, 'r') as zin:
        contents = {name: zin.read(name) for name in zin.namelist()}
    ct = contents['[Content_Types].xml'].decode('utf-8')
    ct = ct.replace('wordprocessingml.template.main+xml',
                    'wordprocessingml.document.main+xml')
    contents['[Content_Types].xml'] = ct.encode('utf-8')
    with zipfile.ZipFile(tmp_tpl, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)
    doc = Document(tmp_tpl)
    os.remove(tmp_tpl)
    return doc
 
 
def insert_text_to_bookmark(doc, bookmark_name, text):
    """
    ブックマーク位置にテキストを挿入する。
    テキストに改行（\\n）が含まれる場合は w:br で改行を挿入する。
 
    注意: addnext は「直後に挿入」のため、複数行を正順で表示するには
    行リストを逆順でループして挿入する必要がある。
    """
    for para in doc.paragraphs:
        for child in para._element:
            if child.tag == qn('w:bookmarkStart'):
                if child.get(qn('w:name')) == bookmark_name:
                    lines = text.split('\n')
                    # 逆順で挿入することで、Word上は正順（上から下）に表示される
                    for idx, line in enumerate(reversed(lines)):
                        run = OxmlElement('w:r')
                        t = OxmlElement('w:t')
                        t.text = line
                        t.set(
                            '{http://www.w3.org/XML/1998/namespace}space',
                            'preserve')
                        run.append(t)
                        child.addnext(run)
                        # 最初の行（逆順なので末尾）以外は改行を先に挿入
                        if idx < len(lines) - 1:
                            br_run = OxmlElement('w:r')
                            br = OxmlElement('w:br')
                            br_run.append(br)
                            child.addnext(br_run)
                    return True
    return False
 
 
def insert_image_to_bookmark(doc, bookmark_name, image_path,
                              width_mm, height_mm):
    """ブックマーク位置に画像を挿入する"""
    for para in doc.paragraphs:
        for child in para._element:
            if child.tag == qn('w:bookmarkStart'):
                if child.get(qn('w:name')) == bookmark_name:
                    run = para.add_run()
                    run.add_picture(
                        image_path,
                        width=Mm(width_mm),
                        height=Mm(height_mm))
                    return True
    return False
 
 
def save_one_doc(row, ad_image_path, template_file=None):
    """
    1社分のWordファイルを output/ フォルダに保存する。
    ファイル名: {ID}_{会社名}.docx
    戻り値: 保存したファイルパス
 
    bm_number に挿入する内容:
      ① A列 ID
      ② C列 会社・店名・氏名 ＋「様」
      ③ G列 25年寄付金額（カンマ区切り・円付き）
    """
    company_id   = str(row.iloc[0])            # A列：ID
    company_name = str(row.iloc[2])            # C列：会社・店名・氏名
    amount       = int(row['_金額'])
    amount_25    = int(row.get('_25寄付金', 0))
    w_mm, h_mm   = get_ad_size(amount)
 
    # bm_number に差し込む文字列を組み立て
    name_sama    = f"{company_name} 様"
    amount_str   = f"{amount_25:,}円" if amount_25 > 0 else ""
    bm_number_text = "\n".join(filter(None, [
        company_id,
        name_sama,
        amount_str,
    ]))
 
    safe_name = "".join(
        c for c in company_name if c not in r'\/:*?"<>|　')
    out_filename = f"{company_id}_{safe_name}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_filename)
 
    doc = open_template(template_file)
    insert_text_to_bookmark(doc, 'bm_number', bm_number_text)
    if ad_image_path and os.path.exists(ad_image_path):
        ok = insert_image_to_bookmark(
            doc, 'bm_AD', ad_image_path, w_mm, h_mm)
        if not ok:
            print(f"  ⚠ ブックマーク bm_AD が見つかりませんでした。")
 
    doc.save(out_path)
    return out_path
 
 
# ============================================================
# GUI: 広告確認・範囲選択ウィンドウ（ページ送り機能付き）
# ============================================================
 
_GUI_RESULT = [None]
 
 
class AdSelectorGUI:
    """
    PDFページを表示して広告範囲をマウスで選択させるGUI。
    ◀前ページ / 次ページ▶ ボタンでページを手動移動可能。
    """
 
    def __init__(self, root, page_images, start_page,
                 company_name, amount, index, total):
        _GUI_RESULT[0] = None
        self.root        = root
        self.page_images = page_images                         # {ページ番号: PIL.Image}
        self.page_list   = sorted(page_images.keys())         # ページ番号のリスト
        self.company_name = company_name
        self.amount      = amount
        self.index       = index
        self.total       = total
 
        # 表示中ページのインデックス（page_list内の位置）
        if start_page in self.page_list:
            self.current_idx = self.page_list.index(start_page)
        else:
            self.current_idx = 0  # 見つからなければ先頭から
 
        # 選択座標
        self.selecting = False
        self.start_x = self.start_y = 0
        self.end_x   = self.end_y   = 0
        self.rect_id = None
        self.scale   = None  # 初回update_canvas()で計算・以降固定
 
        # ── レイアウト構築 ──────────────────────
        self.root.title(
            f"広告確認 [{index}/{total}]: {company_name}　"
            f"({APP_VERSION})")
 
        # 情報ラベル
        w_mm, h_mm = get_ad_size(amount)
        tk.Label(root,
                 text=f"会社名: {company_name}　"
                      f"金額: {amount:,}円　"
                      f"期待サイズ: 幅{w_mm}mm × 高さ{h_mm}mm",
                 font=("Meiryo", 11, "bold"), fg="navy").pack(pady=4)
 
        # バージョン表示（画面右上）
        tk.Label(root, text=APP_VERSION,
                 font=("Meiryo", 8), fg="gray").place(
                     relx=1.0, x=-6, y=4, anchor="ne")
 
        # ── ページ送りコントロール（上段）──────────
        nav_frame = tk.Frame(root)
        nav_frame.pack(pady=(4, 0))
        tk.Button(nav_frame, text="◀ 前ページ",
                  font=("Meiryo", 10), width=12,
                  command=self.prev_page).pack(side=tk.LEFT, padx=4)
        self.page_label = tk.Label(
            nav_frame, text="", font=("Meiryo", 10, "bold"),
            fg="darkred", width=18)
        self.page_label.pack(side=tk.LEFT)
        tk.Button(nav_frame, text="次ページ ▶",
                  font=("Meiryo", 10), width=12,
                  command=self.next_page).pack(side=tk.LEFT, padx=4)
 
        # ── ページ番号直接入力（下段）──────────────
        jump_frame = tk.Frame(root)
        jump_frame.pack(pady=(2, 0))
        tk.Label(jump_frame, text="ページ番号を入力:",
                 font=("Meiryo", 10)).pack(side=tk.LEFT, padx=(0, 4))
        self.page_entry = tk.Entry(
            jump_frame, font=("Meiryo", 11, "bold"),
            width=6, justify="center")
        self.page_entry.pack(side=tk.LEFT)
        self.page_entry.bind("<Return>", lambda e: self.jump_to_page())
        tk.Button(jump_frame, text="ジャンプ",
                  font=("Meiryo", 10), width=8, bg="#2980b9", fg="white",
                  command=self.jump_to_page).pack(side=tk.LEFT, padx=4)
        tk.Label(jump_frame,
                 text=f"（範囲: {self.page_list[0]} 〜 {self.page_list[-1]}）",
                 font=("Meiryo", 9), fg="gray").pack(side=tk.LEFT, padx=4)
 
        tk.Label(root,
                 text="【操作】広告をマウスでドラッグして選択"
                      " → 「✅ この範囲でOK」を押す",
                 font=("Meiryo", 9), fg="dimgray").pack(pady=(4, 2))
 
        # キャンバス（サイズは後で update_canvas() で設定）
        self.canvas = tk.Canvas(root, cursor="crosshair",
                                bg="#cccccc")
        self.canvas.pack(pady=4)
        self.canvas.bind("<ButtonPress-1>",   self.on_press)
        self.canvas.bind("<B1-Motion>",       self.on_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_release)
 
        # アクションボタン
        btn_frame = tk.Frame(root)
        btn_frame.pack(pady=6)
        tk.Button(btn_frame, text="✅ この範囲でOK",
                  font=("Meiryo", 11, "bold"),
                  bg="#27ae60", fg="white", width=18,
                  command=self.cmd_ok).pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text="⏭ この会社をスキップ",
                  font=("Meiryo", 11),
                  bg="#e67e22", fg="white", width=18,
                  command=self.cmd_skip).pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text="💾 中断して保存",
                  font=("Meiryo", 11),
                  bg="#7f8c8d", fg="white", width=18,
                  command=self.cmd_stop).pack(side=tk.LEFT, padx=8)
 
        # 初回描画
        self.update_canvas()
 
    # ── ページ操作 ──────────────────────────────
 
    def current_page_num(self):
        return self.page_list[self.current_idx]
 
    def prev_page(self):
        if self.current_idx > 0:
            self.current_idx -= 1
            self.clear_selection()
            self.update_canvas()
 
    def next_page(self):
        if self.current_idx < len(self.page_list) - 1:
            self.current_idx += 1
            self.clear_selection()
            self.update_canvas()
 
    def jump_to_page(self):
        """入力ボックスのページ番号に直接ジャンプする"""
        val = self.page_entry.get().strip()
        try:
            target = int(val)
        except ValueError:
            messagebox.showwarning("入力エラー",
                                   f"「{val}」は数字ではありません。")
            return
        if target not in self.page_list:
            lo, hi = self.page_list[0], self.page_list[-1]
            messagebox.showwarning("範囲外",
                                   f"ページ {target} は存在しません。\n"
                                   f"有効範囲: {lo} 〜 {hi}")
            return
        self.current_idx = self.page_list.index(target)
        self.clear_selection()
        self.update_canvas()
        self.page_entry.delete(0, tk.END)  # 入力欄をクリア
 
    def clear_selection(self):
        self.start_x = self.start_y = 0
        self.end_x   = self.end_y   = 0
        if self.rect_id:
            self.canvas.delete(self.rect_id)
            self.rect_id = None
 
    def update_canvas(self):
        """現在のページ画像をキャンバスに描画する"""
        page_num   = self.current_page_num()
        page_image = self.page_images[page_num]
 
        # ページラベル更新
        total_pages = len(self.page_list)
        self.page_label.config(
            text=f"ページ {page_num}"
                 f" ({self.current_idx + 1}/{total_pages})")
 
        # スケールは初回のみ計算して固定（ページ切替でウィンドウ位置が動くのを防止）
        if self.scale is None:
            screen_h   = self.root.winfo_screenheight() - 380
            self.scale = screen_h / page_image.height
 
        disp_w = int(page_image.width  * self.scale)
        disp_h = int(page_image.height * self.scale)
        disp_img = page_image.resize((disp_w, disp_h), Image.LANCZOS)
        self.tk_image = ImageTk.PhotoImage(disp_img)
 
        self.canvas.config(width=disp_w, height=disp_h)
        self.canvas.delete("all")
        self.canvas.create_image(0, 0, anchor=tk.NW, image=self.tk_image)
 
    # ── マウス操作 ──────────────────────────────
 
    def on_press(self, event):
        self.selecting = True
        self.start_x, self.start_y = event.x, event.y
        # 既存の矩形があれば非表示にするだけ（削除しない）
        if self.rect_id:
            self.canvas.delete(self.rect_id)
            self.rect_id = None
 
    def on_drag(self, event):
        if self.selecting:
            if self.rect_id:
                # 削除せず座標だけ更新（描画が軽い）
                self.canvas.coords(
                    self.rect_id,
                    self.start_x, self.start_y, event.x, event.y)
            else:
                self.rect_id = self.canvas.create_rectangle(
                    self.start_x, self.start_y, event.x, event.y,
                    outline="red", width=2, dash=(4, 2))
 
    def on_release(self, event):
        self.selecting = False
        self.end_x, self.end_y = event.x, event.y
        # 確定した矩形を実線に変更して視認しやすくする
        if self.rect_id:
            self.canvas.itemconfig(self.rect_id, dash=(), width=3)
 
    # ── ボタン ──────────────────────────────────
 
    def cmd_ok(self):
        x1 = int(min(self.start_x, self.end_x) / self.scale)
        y1 = int(min(self.start_y, self.end_y) / self.scale)
        x2 = int(max(self.start_x, self.end_x) / self.scale)
        y2 = int(max(self.start_y, self.end_y) / self.scale)
        if x2 - x1 < 10 or y2 - y1 < 10:
            messagebox.showwarning(
                "選択エラー",
                "範囲が小さすぎます。\nもう一度ドラッグしてください。")
            return
 
        page_num   = self.current_page_num()
        page_image = self.page_images[page_num]
 
        # ── 自動改ページ判定 ──────────────────────
        # ※ update_canvas() を呼ぶ前に判定すること
        #   （update_canvas後はself.scaleが変わるためy2が狂う）
        page_h_px        = page_image.height        # 原寸ページ高さ（px）
        page_w_px        = page_image.width         # 原寸ページ幅（px）
        PAGE_H_MM        = 265.0                    # B5変型判の高さ（mm）
        PAGE_W_MM        = 182.0                    # B5変型判の幅（mm）
        px_per_mm_h      = page_h_px / PAGE_H_MM
        px_per_mm_w      = page_w_px / PAGE_W_MM
        bottom_margin_mm = (page_h_px - y2) / px_per_mm_h  # 下端からの距離
        right_margin_mm  = (page_w_px - x2) / px_per_mm_w  # 右端からの距離
 
        # 下端30mm以内 かつ 右端30mm以内 → 改ページ
        auto_next = (bottom_margin_mm <= 30.0) and (right_margin_mm <= 30.0)
 
        # デバッグ用：コンソールに距離を表示
        print(f"  　下端: {bottom_margin_mm:.1f}mm　右端: {right_margin_mm:.1f}mm"
              f" → {'次ページへ自動移動' if auto_next else '同ページ継続'}")
 
        # 結果を確定（auto_nextフラグも返す）
        _GUI_RESULT[0] = ("ok", page_num, x1, y1, x2 - x1, y2 - y1,
                          auto_next)
 
        # 改ページ判定後にupdate_canvas（scaleが変わってもOK）
        if auto_next and self.current_idx < len(self.page_list) - 1:
            self.current_idx += 1
            self.clear_selection()
            self.update_canvas()
            self.page_label.config(
                text=f"✅ 保存 → ページ {self.current_page_num()} へ",
                fg="green")
 
        self.root.destroy()
 
    def cmd_skip(self):
        _GUI_RESULT[0] = ("skip",)
        self.root.destroy()
 
    def cmd_stop(self):
        _GUI_RESULT[0] = ("stop",)
        self.root.destroy()
 
 
def select_ad_region(tk_root, page_images, start_page,
                     company_name, amount, index, total):
    """GUIを表示して広告範囲を選択させる（Toplevel を使用）"""
    # tk.Tk() は main() で1回だけ作成済み。ここでは Toplevel を使う。
    win = tk.Toplevel(tk_root)
    win.resizable(False, False)  # リサイズ不可にして位置ずれを防止
    AdSelectorGUI(win, page_images, start_page,
                  company_name, amount, index, total)
    # ウィンドウサイズ確定後に画面中央へ配置
    win.update_idletasks()
    sw = tk_root.winfo_screenwidth()
    sh = tk_root.winfo_screenheight()
    ww = win.winfo_width()
    wh = win.winfo_height()
    x  = (sw - ww) // 2
    y  = max(0, (sh - wh) // 2 - 20)
    win.geometry(f"+{x}+{y}")
    win.grab_set()          # このウィンドウをモーダルにする
    tk_root.wait_window(win)  # ウィンドウが閉じるまで待機
    result = _GUI_RESULT[0]
    return result if result else ("skip",)
 
 
# ============================================================
# メイン処理
# ============================================================
 
def select_files_dialog(root, last_dir=None):
    """
    起動時にExcel・PDF・Wordテンプレートのファイルを選択するダイアログを表示する。
    root    : メインの tk.Tk() ウィンドウ（呼び出し元で作成済み）
    last_dir: 前回記憶したフォルダ（progress.jsonから復元）
    戻り値: (excel_path, pdf_path, template_path, excel_dir) のタプル
            キャンセルされた場合は None を返す。
    """
    # 初期フォルダ：前回記憶フォルダ → DATA_DIR の順で使用
    init_dir = last_dir if (last_dir and os.path.isdir(last_dir)) else DATA_DIR

    # ── Excel選択 ──────────────────────────────
    messagebox.showinfo(
        "ファイル選択 (1/3)",
        "① Excelファイル（.xlsx）を選択してください。",
        parent=root)
    excel_path = filedialog.askopenfilename(
        title="Excelファイルを選択",
        filetypes=[("Excelファイル", "*.xlsx")],
        initialdir=init_dir,
        parent=root)
    if not excel_path:
        messagebox.showwarning("キャンセル", "Excelファイルが選択されませんでした。終了します。",
                               parent=root)
        return None

    # Excelのフォルダを以降のダイアログのデフォルトに使用
    excel_dir = os.path.dirname(excel_path)

    # ── PDF選択 ────────────────────────────────
    messagebox.showinfo(
        "ファイル選択 (2/3)",
        "② PDFファイル（.pdf）を選択してください。",
        parent=root)
    pdf_path = filedialog.askopenfilename(
        title="PDFファイルを選択",
        filetypes=[("PDFファイル", "*.pdf")],
        initialdir=excel_dir,
        parent=root)
    if not pdf_path:
        messagebox.showwarning("キャンセル", "PDFファイルが選択されませんでした。終了します。",
                               parent=root)
        return None

    # ── Wordテンプレート選択 ────────────────────
    messagebox.showinfo(
        "ファイル選択 (3/3)",
        "③ Wordテンプレート（.dotx）を選択してください。",
        parent=root)
    template_path = filedialog.askopenfilename(
        title="Wordテンプレートを選択",
        filetypes=[("Wordテンプレート", "*.dotx")],
        initialdir=excel_dir,
        parent=root)
    if not template_path:
        messagebox.showwarning("キャンセル", "Wordテンプレートが選択されませんでした。終了します。",
                               parent=root)
        return None

    return excel_path, pdf_path, template_path, excel_dir
 
 
def main():
    print("=" * 55)
    print("  寄付依頼書 自動生成ツール  v2.7")
    print("=" * 55)
 
    # フォルダ作成（TMP_DIRのみ。OUTPUT_DIRはExcel選択後に決定）
    os.makedirs(TMP_DIR, exist_ok=True)
 
    # 進捗読み込み（ファイル選択前に行い、前回フォルダを取得）
    progress    = load_progress()
    last_dir    = progress.get("last_dir", None)
 
    # ── Tk ルートウィンドウを1回だけ作成（EXE対応）──
    _tk_root = tk.Tk()
    _tk_root.withdraw()   # 本体ウィンドウは非表示（ダイアログ・GUIはToplevelで出す）

    # ── ファイル選択ダイアログ ──────────────────
    result = select_files_dialog(_tk_root, last_dir)
    if result is None:
        _tk_root.destroy()
        return
    excel_file, pdf_file, template_file, excel_dir = result
    print(f"\n  Excel    : {os.path.basename(excel_file)}")
    print(f"  PDF      : {os.path.basename(pdf_file)}")
    print(f"  テンプレ : {os.path.basename(template_file)}")
 
    # output フォルダをExcelと同じフォルダ内に設定
    global OUTPUT_DIR
    OUTPUT_DIR = os.path.join(excel_dir, "output")
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"  出力先   : {OUTPUT_DIR}")
 
    # Excelデータ読み込み
    print("\nExcelを読み込み中...")
    df    = load_excel(excel_file)
    total = len(df)
 
    # 進捗の続き情報を取得
    start_index = progress.get("last_index", 0)
    completed   = set(progress.get("completed_ids", []))
 
    # 再開確認：途中保存があり、かつ全件完了でない場合のみ表示
    if 0 < start_index < total:
        resume = messagebox.askyesno(
            "前回の続きがあります",
            f"前回の途中データがあります。\n\n"
            f"  完了済み : {len(completed)} 社 / 全 {total} 社\n"
            f"  次の処理 : {start_index + 1} 件目から\n\n"
            f"続きから再開しますか？\n"
            f"（「いいえ」を選ぶと最初からやり直します）",
            parent=_tk_root)
        if not resume:
            start_index = 0
            completed   = set()
            progress    = {"completed_ids": [], "last_index": 0}
    else:
        # 全件完了済みまたは新規の場合は最初から
        start_index = 0
        completed   = set()
        progress    = {"completed_ids": [], "last_index": 0}
        messagebox.showinfo("処理開始", f"全 {total} 社を処理します。", parent=_tk_root)
 
    # 今回選んだフォルダをprogress.jsonに保存（次回起動時に使用）
    progress["last_dir"] = excel_dir
 
    # PDF → 画像変換のみ（テキスト検索は使用しない）
    print()
    page_images = pdf_to_images(pdf_file)
 
    # メインループ
    print()
    print("-" * 55)
    # 再開時は保存されたページから、新規は None
    prev_page = progress.get("last_page", None) if start_index > 0 else None
    for i in range(start_index, total):
        row          = df.iloc[i]
        # A列=ID(0列目)、C列=会社名(2列目) を列位置で取得
        company_id   = str(row.iloc[0])            # A列：ID
        company_name = str(row.iloc[2])            # C列：会社・店名・氏名
        amount       = int(row['_金額'])
        f_page       = int(row.get('_page', 0))   # X列：プログラムページ
 
        # ── 表示開始ページの決定 ──────────────────
        # ① F列に番号あり → そのページを開く
        # ② F列が空（0）→ 直前OKページの次ページ、なければ先頭
        all_pages = sorted(page_images.keys())
        if f_page > 0 and f_page in page_images:
            found_page = f_page
            hint = f"F列指定 ページ {found_page}"
        elif prev_page is not None:
            # prev_pageをそのまま使う
            # （prev_pageはOK時に自動改ページ有無に応じて正しく更新済み）
            found_page = prev_page
            hint = f"順番 ページ {found_page}（前社p{prev_page}から）"
        else:
            found_page = all_pages[0]
            hint = f"先頭ページ {found_page}"
 
        print(f"\n[{i+1}/{total}] {company_name}"
              f"  (ID:{company_id} / {amount:,}円)"
              f"\n  → 📄 {hint}")
 
        # GUI表示して広告範囲を選択
        result = select_ad_region(
            _tk_root, page_images, found_page,
            company_name, amount, i + 1, total)
        action = result[0]
 
        # ── 中断 ────────────────────────────────
        if action == "stop":
            print("\n⏹ 処理を中断します...")
            progress["last_index"]    = i
            progress["completed_ids"] = list(completed)
            progress["last_page"]     = prev_page  # 中断時のページを保存
            save_progress(progress)
            print(f"  完了済み: {len(completed)} 社")
            print(f"  次回は [{i+1}/{total}] "
                  f"{company_name} から再開します。")
            messagebox.showinfo("中断しました",
                f"処理を中断しました。\n\n"
                f"  完了済み: {len(completed)} 社\n"
                f"  次回は [{i+1}/{total}] {company_name} から再開します。",
                parent=_tk_root)
            _tk_root.destroy()
            return
 
        # ── スキップ ────────────────────────────
        elif action == "skip":
            print(f"  → スキップしました。")
            continue
 
        # ── OK: Wordファイル生成 ─────────────────
        elif action == "ok":
            _, page_num, x, y, w, h, auto_next = result
 
            # prev_page の更新：
            #   自動改ページあり → GUIはすでに次ページを表示済み
            #                      → 次の会社は同じページから始める
            #                      → prev_page = 次ページ番号
            #   自動改ページなし → 同ページに次の会社が続く可能性あり
            #                      → prev_page = 現ページ番号（進めない）
            all_pages = sorted(page_images.keys())
            if auto_next and page_num in all_pages:
                next_idx = all_pages.index(page_num) + 1
                if next_idx < len(all_pages):
                    prev_page = all_pages[next_idx]  # 次ページを引き継ぐ
                else:
                    prev_page = page_num
            else:
                prev_page = page_num  # 同ページ継続
 
            # 広告を切り出してグレースケールPNG保存
            ad_img  = page_images[page_num].crop(
                (x, y, x + w, y + h))
            ad_gray = ad_img.convert('L')
            tmp_png = os.path.join(TMP_DIR, f"ad_{company_id}.png")
            ad_gray.save(tmp_png, 'PNG')
 
            # 1社分のWordファイルを生成・保存
            out_path = save_one_doc(row, tmp_png, template_file)
            completed.add(company_id)
            print(f"  → 保存しました: "
                  f"{os.path.basename(out_path)}")
 
            # 10社ごとに自動で進捗保存
            if len(completed) % 10 == 0:
                progress["last_index"]    = i + 1
                progress["completed_ids"] = list(completed)
                save_progress(progress)
 
    # ── 全件完了 ─────────────────────────────────
    progress["last_index"]    = total
    progress["completed_ids"] = list(completed)
    save_progress(progress)
 
    print()
    print("=" * 55)
    print(f"✅ 全件処理完了！")
    print(f"   生成ファイル数 : {len(completed)} 社")
    print(f"   出力フォルダ   : {OUTPUT_DIR}")
    print("=" * 55)
    messagebox.showinfo("完了",
        f"全件処理完了！\n\n"
        f"  生成ファイル数 : {len(completed)} 社\n"
        f"  出力フォルダ   : {OUTPUT_DIR}",
        parent=_tk_root)
    _tk_root.destroy()
 
 
if __name__ == "__main__":
    main()